"""Generate export artefacts (Word / PDF / Excel / CSV / Markdown) from analysis output."""
from __future__ import annotations

import csv
import io
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from config import settings

NAVY = colors.HexColor("#0d2240")
GOLD = colors.HexColor("#f5a800")
INK = colors.HexColor("#1a2d44")
MUTED = colors.HexColor("#6b8299")
LINE = colors.HexColor("#e2eaf2")

EXPORT_DIR = Path(settings.upload_dir) / "exports"
EXPORT_DIR.mkdir(parents=True, exist_ok=True)


def _split_blocks(md: str) -> list[tuple[str, str]]:
    """Yield (kind, text) blocks where kind is 'h1' 'h2' 'h3' 'table' 'li' 'p'."""
    blocks: list[tuple[str, str]] = []
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1
        elif line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
            i += 1
        elif line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            tbl = [line]
            i += 2  # skip separator
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i])
                i += 1
            blocks.append(("table", "\n".join(tbl)))
        elif re.match(r"^\s*[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]))
                i += 1
            blocks.append(("li", "\n".join(items)))
        elif re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i]))
                i += 1
            blocks.append(("ol", "\n".join(items)))
        else:
            para = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", "|", "- ", "* ")):
                para.append(lines[i])
                i += 1
            blocks.append(("p", " ".join(para)))
    return blocks


def _table_rows(tbl_md: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in tbl_md.splitlines():
        if not line.strip().startswith("|"):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def _strip_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


# ---------- MARKDOWN ----------
def export_markdown(content: str, meta: dict) -> str:
    fname = f"{meta['id']}.md"
    path = EXPORT_DIR / fname
    header = (
        f"# {meta['mode_label']}\n"
        f"**Project:** {meta.get('project_name', 'Quick Analysis')}  \n"
        f"**Generated:** {meta.get('completed_at', '')}  \n"
        f"**Model:** {meta.get('model_used', '')}  \n"
        f"**Hash:** `{meta.get('blockchain_hash', '')}`\n\n---\n\n"
    )
    path.write_text(header + content, encoding="utf-8")
    return str(path)


# ---------- CSV ----------
def export_csv(content: str, meta: dict) -> str:
    fname = f"{meta['id']}.csv"
    path = EXPORT_DIR / fname
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["section", "type", "content"])
    for kind, text in _split_blocks(content):
        if kind == "table":
            for row in _table_rows(text):
                w.writerow(["table-row", "row", " | ".join(row)])
        else:
            w.writerow([kind, kind, _strip_md(text).replace("\n", " ")])
    path.write_text(buf.getvalue(), encoding="utf-8")
    return str(path)


# ---------- EXCEL ----------
def export_xlsx(content: str, meta: dict) -> str:
    fname = f"{meta['id']}.xlsx"
    path = EXPORT_DIR / fname
    wb = Workbook()
    ws = wb.active
    ws.title = "Report"
    header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=12)
    header_fill = PatternFill("solid", fgColor="0D2240")

    ws["A1"] = meta["mode_label"]
    ws["A1"].font = Font(name="Calibri", bold=True, size=16, color="0D2240")
    ws["A2"] = f"Project: {meta.get('project_name', 'Quick Analysis')}"
    ws["A3"] = f"Generated: {meta.get('completed_at', '')}"
    ws["A4"] = f"Model: {meta.get('model_used', '')}"
    ws["A5"] = f"Hash: {meta.get('blockchain_hash', '')}"
    ws.column_dimensions["A"].width = 40

    row = 7
    for kind, text in _split_blocks(content):
        if kind in ("h1", "h2", "h3"):
            ws.cell(row=row, column=1, value=_strip_md(text)).font = Font(
                bold=True, size=13 if kind == "h1" else 12, color="0D2240"
            )
            row += 1
        elif kind == "table":
            rows = _table_rows(text)
            if not rows:
                continue
            header = rows[0]
            for c, h in enumerate(header, start=1):
                cell = ws.cell(row=row, column=c, value=_strip_md(h))
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center")
            row += 1
            for r in rows[1:]:
                for c, v in enumerate(r, start=1):
                    ws.cell(row=row, column=c, value=_strip_md(v))
                row += 1
            row += 1
        elif kind in ("li", "ol"):
            for it in text.splitlines():
                ws.cell(row=row, column=1, value="• " + _strip_md(it))
                row += 1
        else:
            ws.cell(row=row, column=1, value=_strip_md(text))
            row += 1

    wb.save(str(path))
    return str(path)


# ---------- WORD ----------
def export_docx(content: str, meta: dict) -> str:
    fname = f"{meta['id']}.docx"
    path = EXPORT_DIR / fname
    doc = Document()

    # Cover
    doc.add_paragraph("4XSTRUCT · STRUCTMIND AI").runs[0].font.size = Pt(10)
    title = doc.add_paragraph()
    tr = title.add_run(meta["mode_label"].upper())
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = RGBColor(0x0D, 0x22, 0x40)
    doc.add_paragraph(f"Project: {meta.get('project_name', 'Quick Analysis')}")
    doc.add_paragraph(f"Generated: {meta.get('completed_at', '')}")
    doc.add_paragraph(f"Model: {meta.get('model_used', '')}")
    doc.add_paragraph(f"SHA-256 Hash: {meta.get('blockchain_hash', '')}")
    doc.add_paragraph("")

    for kind, text in _split_blocks(content):
        if kind in ("h1", "h2", "h3"):
            lvl = {"h1": 1, "h2": 2, "h3": 3}[kind]
            h = doc.add_heading(_strip_md(text), level=lvl)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x0D, 0x22, 0x40)
        elif kind == "table":
            rows = _table_rows(text)
            if not rows:
                continue
            tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
            tbl.style = "Light Grid Accent 1"
            for ri, r in enumerate(rows):
                for ci, v in enumerate(r):
                    tbl.cell(ri, ci).text = _strip_md(v)
            doc.add_paragraph("")
        elif kind in ("li", "ol"):
            for it in text.splitlines():
                doc.add_paragraph(_strip_md(it), style="List Bullet" if kind == "li" else "List Number")
        else:
            doc.add_paragraph(_strip_md(text))

    doc.save(str(path))
    return str(path)


# ---------- PDF ----------
def export_pdf(content: str, meta: dict) -> str:
    fname = f"{meta['id']}.pdf"
    path = EXPORT_DIR / fname
    styles = getSampleStyleSheet()

    story = []
    brand = ParagraphStyle(
        "brand",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        textColor=GOLD,
        fontSize=10,
        spaceAfter=6,
    )
    h1 = ParagraphStyle(
        "h1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        textColor=NAVY,
        fontSize=22,
        spaceAfter=14,
    )
    h2 = ParagraphStyle(
        "h2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        textColor=NAVY,
        fontSize=16,
        spaceAfter=10,
    )
    h3 = ParagraphStyle(
        "h3",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        textColor=NAVY,
        fontSize=13,
        spaceAfter=8,
    )
    body = ParagraphStyle(
        "body",
        parent=styles["Normal"],
        fontName="Helvetica",
        textColor=INK,
        fontSize=10.5,
        leading=15,
        spaceAfter=8,
    )
    meta_s = ParagraphStyle(
        "meta", parent=body, textColor=MUTED, fontSize=9, spaceAfter=4
    )

    story.append(Paragraph("4XSTRUCT · STRUCTMIND AI", brand))
    story.append(Paragraph(meta["mode_label"].upper(), h1))
    story.append(Paragraph(f"Project: {meta.get('project_name', 'Quick Analysis')}", meta_s))
    story.append(Paragraph(f"Generated: {meta.get('completed_at', '')}", meta_s))
    story.append(Paragraph(f"Model: {meta.get('model_used', '')}", meta_s))
    story.append(Paragraph(f"SHA-256: {meta.get('blockchain_hash', '')}", meta_s))
    story.append(Spacer(1, 0.2 * inch))

    for kind, text in _split_blocks(content):
        clean = _strip_md(text)
        if kind == "h1":
            story.append(Paragraph(clean, h1))
        elif kind == "h2":
            story.append(Paragraph(clean, h2))
        elif kind == "h3":
            story.append(Paragraph(clean, h3))
        elif kind == "table":
            rows = _table_rows(text)
            if not rows:
                continue
            data = [[Paragraph(_strip_md(c), body) for c in r] for r in rows]
            tbl = Table(data, repeatRows=1)
            tbl.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("GRID", (0, 0), (-1, -1), 0.4, LINE),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f9fc")]),
                    ]
                )
            )
            story.append(tbl)
            story.append(Spacer(1, 0.15 * inch))
        elif kind in ("li", "ol"):
            for idx, it in enumerate(text.splitlines(), 1):
                bullet = "•" if kind == "li" else f"{idx}."
                story.append(Paragraph(f"{bullet}  {_strip_md(it)}", body))
        else:
            story.append(Paragraph(clean, body))

    pdf = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.85 * inch,
        bottomMargin=0.75 * inch,
        title=meta["mode_label"],
        author="StructMind AI",
    )
    pdf.build(story)
    return str(path)


EXPORTERS = {
    "markdown": export_markdown,
    "csv": export_csv,
    "xlsx": export_xlsx,
    "docx": export_docx,
    "pdf": export_pdf,
}


def generate_all_exports(content: str, meta: dict) -> list[dict]:
    results = []
    for fmt, fn in EXPORTERS.items():
        try:
            p = fn(content, meta)
            results.append({"format": fmt, "path": p})
        except Exception as e:  # noqa: BLE001
            results.append({"format": fmt, "path": "", "error": str(e)})
    return results
